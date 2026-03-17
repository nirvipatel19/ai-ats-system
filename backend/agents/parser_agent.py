import logging
import pdfplumber
from docx import Document

logger = logging.getLogger(__name__)


def parse_docx_text(docx_content: bytes) -> tuple[str, bool]:
    """Extract text from Word document. Returns (text, success)."""
    import io
    try:
        doc = Document(io.BytesIO(docx_content))
        parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        full_text = "\n".join(parts)
        return full_text, len(full_text) > 0
    except Exception as e:
        logger.error(f"DOCX parse failed: {e}")
        return "", False


def parse_pdf_text(pdf_content: bytes) -> tuple[str, bool]:
    """Extract text from any PDF using text extraction and OCR fallback. Returns (text, success)."""
    import io
    try:
        text_parts = []
        with pdfplumber.open(io.BytesIO(pdf_content)) as pdf:
            # 1. Try standard text extraction
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            
            # 2. OCR Fallback if text is minimal (e.g., scanned PDF or image-heavy)
            # Threshold: if we got less than 100 characters, try OCR
            current_text = "\n".join(text_parts).strip()
            if len(current_text) < 100:
                logger.info("Normal PDF extraction yielded little text. Attempting OCR fallback...")
                ocr_parts = []
                for page in pdf.pages:
                    try:
                        # Extract text from this specific page to see if it's empty
                        p_text = page.extract_text() or ""
                        # If page is mostly empty/scanned, OCR it
                        if len(p_text.strip()) < 50:
                            import pytesseract
                            from PIL import Image
                            p_img = page.to_image(resolution=300)
                            pil_img = p_img.original
                            ocr_text = pytesseract.image_to_string(pil_img)
                            if ocr_text.strip():
                                ocr_parts.append(ocr_text.strip())
                    except Exception as ocr_e:
                        logger.error(f"OCR failed for a page: {ocr_e}")
                
                if ocr_parts:
                    # Add OCR results as new lines
                    text_parts.extend(ocr_parts)

        full_text = "\n".join(text_parts) if text_parts else ""
        return full_text, len(full_text.strip()) > 0
    except Exception as e:
        logger.error(f"PDF parse failed: {e}")
        return "", False


def parse_document_text(content: bytes, filename: str) -> tuple[str, bool]:
    """Extract text from PDF or DOCX based on filename. Returns (text, success)."""
    fn = (filename or "").lower()
    if fn.endswith(".docx"):
        return parse_docx_text(content)
    return parse_pdf_text(content)


def parse_resume(pdf_content: bytes, filename: str) -> dict:
    """
    Extract text from PDF resume.
    Returns dict with filename and extracted text.
    """
    logger.info("Running Resume Parser Agent...")
    try:
        text, ok = parse_pdf_text(pdf_content)
        full_text = text if ok else ""
        
        # Extract name from first line if possible (heuristic)
        name = filename.replace(".pdf", "").replace("_", " ").replace("-", " ")
        
        result = {
            "filename": filename,
            "name": name,
            "raw_text": full_text,
            "success": len(full_text.strip()) > 0,
        }
        logger.info(f"Resume Parser Agent complete. Extracted {len(full_text)} chars from {filename}")
        return result
    except Exception as e:
        logger.error(f"Resume Parser Agent failed for {filename}: {e}")
        return {
            "filename": filename,
            "name": filename.replace(".pdf", ""),
            "raw_text": "",
            "success": False,
            "error": str(e),
        }
